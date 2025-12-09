import json
import os
import io
import re
import time
import concurrent.futures
from datetime import datetime
from dotenv import load_dotenv
# from openai import OpenAI
import openai
import docx
import fitz  # PyMuPDF
from copy import deepcopy
from docx2pdf import convert

# Load environment variables
load_dotenv()

# Initialize OpenAI client
api_key = os.getenv('AZURE_OPENAI_API_KEY')
if not api_key:
    raise ValueError("AZURE_OPENAI_API_KEY not found in environment. Please set it in your .env file.")

# from  llm import get_llm
# client = get_llm()
import openai

AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
AZURE_OPENAI_VERSION = os.getenv("OPENAI_API_VERSION", "2023-07-01-preview")


client = openai.AzureOpenAI(
    api_version=AZURE_OPENAI_VERSION,
    azure_endpoint=AZURE_OPENAI_ENDPOINT,
    azure_deployment=AZURE_OPENAI_DEPLOYMENT, 
    api_key=AZURE_OPENAI_API_KEY
)

# ============================================================================
# Document Extraction Functions
# ============================================================================

def extract_text_from_docx(file_path):
    """
    Extract text from a Word document including comprehensive table data.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content as string
    """
    try:
        doc = docx.Document(file_path)
        content_parts = []
        last_question = None
        
        def extract_formatted_text(paragraph):
            """Extract text with formatting preserved as HTML-like tags."""
            formatted_text = ""
            for run in paragraph.runs:
                text = run.text
                if not text:
                    continue
                    
                # Apply formatting tags
                if run.bold:
                    text = f"<b>{text}</b>"
                if run.italic:
                    text = f"<i>{text}</i>"
                if run.underline:
                    text = f"<u>{text}</u>"
                if run.strike:
                    text = f"<s>{text}</s>"
                
                # Handle font size
                if hasattr(run, 'font') and run.font.size:
                    size_pt = run.font.size.pt if run.font.size.pt else 11
                    if size_pt > 12:
                        text = f"<large>{text}</large>"
                    elif size_pt < 10:
                        text = f"<small>{text}</small>"
                
                text = text.replace('\n', '<br/>')
                formatted_text += text
            
            if formatted_text.strip():
                formatted_text += "<br/><br/>"
            
            return formatted_text
        
        for element in doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                for para in doc.paragraphs:
                    if para._element == element and para.text.strip():
                        formatted_question = extract_formatted_text(para)
                        last_question = formatted_question
                        break
            elif element.tag.endswith('tbl'):  # Table
                for table in doc.tables:
                    if table._element == element:
                        # Extract table structure
                        table_data = {
                            'headers': [],
                            'rows': [],
                            'structure': 'table'
                        }
                        
                        if table.rows:
                            header_row = table.rows[0]
                            headers = [cell.text.strip() or "[Empty]" for cell in header_row.cells]
                            table_data['headers'] = headers
                            
                            for row in table.rows[1:]:
                                row_data = [cell.text.strip() or "[Empty]" for cell in row.cells]
                                table_data['rows'].append(row_data)
                        
                        # Format table for processing
                        table_text = "[TABLE START]\n"
                        table_text += f"HEADERS: {' | '.join(table_data['headers'])}\n"
                        for idx, row in enumerate(table_data['rows'], 1):
                            table_text += f"ROW {idx}: {' | '.join(row)}\n"
                        table_text += "[TABLE END]"
                        
                        if last_question:
                            content_parts.append(last_question)
                            last_question = None
                        content_parts.append(table_text)
                        break
        
        if last_question:
            content_parts.append(last_question)
        
        final_text = '\n\n'.join(content_parts)
        final_text = re.sub(r'<br/><br/><br/>+', '<br/><br/>', final_text)
        final_text = re.sub(r'<br/><br/>$', '', final_text)
        
        return final_text
    except Exception as e:
        print(f"Error processing DOCX file: {str(e)}")
        return None


def extract_text_from_txt(file_path):
    """Extract text from a text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text
    except UnicodeDecodeError:
        print("Error: Unable to decode the text file. Please ensure it's a valid text file.")
        return None
    except Exception as e:
        print(f"Error processing text file: {str(e)}")
        return None


def extract_text_from_pdf(file_path):
    """Extract text from a PDF file with table detection and formatting preservation."""
    try:
        doc = fitz.open(file_path)
        text_content = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            page_content = []
            
            # Extract tables with full structure
            try:
                tables = page.find_tables()
                if tables and tables.tables:
                    for table_idx, table in enumerate(tables.tables):
                        table_content = []
                        table_content.append(f"[TABLE {table_idx + 1} START]")
                        
                        extracted_table = table.extract()
                        if extracted_table:
                            # First row as headers
                            if len(extracted_table) > 0:
                                headers = []
                                for cell in extracted_table[0]:
                                    cell_text = str(cell).strip() if cell is not None else "[Empty]"
                                    headers.append(cell_text)
                                table_content.append(f"HEADERS: {' | '.join(headers)}")
                                
                                # Remaining rows as data
                                for row_idx, row in enumerate(extracted_table[1:], 1):
                                    row_data = []
                                    for cell in row:
                                        cell_text = str(cell).strip() if cell is not None else "[Empty]"
                                        row_data.append(cell_text)
                                    table_content.append(f"ROW {row_idx}: {' | '.join(row_data)}")
                        
                        table_content.append(f"[TABLE {table_idx + 1} END]")
                        page_content.append('\n'.join(table_content))
            except Exception as table_error:
                print(f"Warning: Error extracting tables from page {page_num + 1}: {str(table_error)}")
            
            # Extract text with formatting information
            try:
                text_dict = page.get_text("dict")
                formatted_text = ""
                
                for block in text_dict.get("blocks", []):
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line.get("spans", []):
                                text = span.get("text", "")
                                if not text:
                                    continue
                                
                                # Check for formatting based on font properties
                                font_flags = span.get("flags", 0)
                                font_size = span.get("size", 11)
                                
                                # Apply formatting tags
                                if font_flags & 16:  # Bold
                                    text = f"<b>{text}</b>"
                                if font_flags & 2:   # Italic
                                    text = f"<i>{text}</i>"
                                if font_flags & 32:  # Underline
                                    text = f"<u>{text}</u>"
                                if font_flags & 64:  # Strikeout
                                    text = f"<s>{text}</s>"
                                
                                # Handle font size
                                if font_size > 12:
                                    text = f"<large>{text}</large>"
                                elif font_size < 10:
                                    text = f"<small>{text}</small>"
                                
                                formatted_text += text
                            
                            formatted_text += "<br/>"
                        formatted_text += "<br/>"
                
                if formatted_text.strip():
                    page_content.append(formatted_text.strip())
                    
            except Exception as text_error:
                print(f"Warning: Error extracting formatted text from page {page_num + 1}: {str(text_error)}")
                # Fallback to plain text extraction
                try:
                    text = page.get_text("text")
                    if text and text.strip():
                        page_content.append(text.strip())
                except Exception as fallback_error:
                    print(f"Warning: Error extracting plain text from page {page_num + 1}: {str(fallback_error)}")
            
            if page_content:
                text_content.extend(page_content)
        
        doc.close()
        
        if not text_content:
            print("No text content could be extracted from the PDF")
            return None
            
        return '\n\n'.join(text_content)
    except Exception as e:
        print(f"Error processing PDF file: {str(e)}")
        return None


def convert_docx_to_pdf(docx_path):
    """
    Convert a DOCX file to PDF format.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Path to the converted PDF file, or None if conversion fails
    """
    try:
        # Generate PDF path in the same directory as the DOCX file
        base_name = os.path.splitext(docx_path)[0]
        pdf_path = f"{base_name}_converted.pdf"
        
        print(f"Converting DOCX to PDF: {os.path.basename(docx_path)} -> {os.path.basename(pdf_path)}")
        convert(docx_path, pdf_path)
        
        if os.path.exists(pdf_path):
            print(f"✓ Successfully converted to PDF: {pdf_path}")
            return pdf_path
        else:
            print("❌ PDF file was not created")
            return None
    except Exception as e:
        print(f"❌ Error converting DOCX to PDF: {str(e)}")
        print("Note: docx2pdf requires Microsoft Word to be installed on Windows.")
        return None


def extract_text_from_document(file_path):
    """
    Extract text from a document based on its file extension.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Extracted text content as string
    """
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.txt'):
        return extract_text_from_txt(file_path)
    elif file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_path}")


# ============================================================================
# Text Processing Functions
# ============================================================================

def split_text_into_chunks(text, max_chars=20000):
    """Split text into chunks based purely on character count for reliable processing."""
    if len(text) <= max_chars:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate end position
        end = start + max_chars
        
        if end >= len(text):
            # Last chunk
            chunks.append(text[start:])
            break
        
        # Try to find a good break point near the end (prefer line breaks)
        chunk_text = text[start:end]
        
        # Look for line breaks in the last 5000 characters to avoid cutting mid-sentence
        search_start = max(0, len(chunk_text) - 5000)
        last_newline = chunk_text.rfind('\n', search_start)
        
        if last_newline > search_start:
            # Found a good break point
            actual_end = start + last_newline + 1
            chunks.append(text[start:actual_end].strip())
            start = actual_end
        else:
            # No good break point found, just cut at character limit
            chunks.append(text[start:end].strip())
            start = end
    
    return [chunk for chunk in chunks if chunk.strip()]


def count_tokens_estimate(text):
    """Rough estimate of token count (approximately 4 characters per token)"""
    return len(text) // 4


# ============================================================================
# GPT-4 Processing Functions
# ============================================================================

def _clean_llm_response_for_json(response_content):
    """Clean LLM response to fix common JSON parsing issues"""
    try:
        # Remove any markdown code blocks
        if '```json' in response_content:
            start = response_content.find('```json') + 7
            end = response_content.find('```', start)
            if end != -1:
                response_content = response_content[start:end].strip()
        elif '```' in response_content:
            start = response_content.find('```') + 3
            end = response_content.find('```', start)
            if end != -1:
                response_content = response_content[start:end].strip()
        
        # Find the first { and last } to extract just the JSON content
        first_brace = response_content.find('{')
        last_brace = response_content.rfind('}')
        
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            response_content = response_content[first_brace:last_brace + 1]
        
        # Fix common JSON issues
        import re
        
        # Remove any trailing commas before closing braces/brackets
        response_content = re.sub(r',(\s*[}\]])', r'\1', response_content)
        
        # Enhanced string fixing for very long responses
        # Look for unterminated strings and try to fix them
        lines = response_content.split('\n')
        fixed_lines = []
        
        for line in lines:
            # Count quotes in the line
            quote_count = line.count('"')
            if quote_count % 2 != 0:  # Odd number of quotes
                # Try to find where the string should end
                if '"text":' in line:
                    # Look for the next quote or end of line
                    text_start = line.find('"text":') + 7
                    remaining = line[text_start:]
                    
                    # Find the next quote or end
                    next_quote = remaining.find('"', 1)
                    if next_quote > 0:
                        # String is properly closed
                        fixed_lines.append(line)
                    else:
                        # String is unterminated, try to close it
                        if line.endswith(','):
                            fixed_lines.append(line[:-1] + '",')
                        else:
                            fixed_lines.append(line + '"')
                else:
                    # Other unterminated string, try to close it
                    if line.endswith(','):
                        fixed_lines.append(line[:-1] + '",')
                    else:
                        fixed_lines.append(line + '"')
            else:
                fixed_lines.append(line)
        
        response_content = '\n'.join(fixed_lines)
        
        # Simple string fixing: look for unterminated strings and close them
        # This is a more conservative approach that doesn't break valid JSON
        if response_content.count('"') % 2 != 0:
            # If we have an odd number of quotes, try to close the last unterminated string
            last_quote = response_content.rfind('"')
            if last_quote > 0:
                # Look for the last field that might be unterminated
                if response_content[last_quote-1] != '"':  # Not a double quote
                    # Try to close the string
                    if response_content.endswith(','):
                        response_content = response_content[:-1] + '",'
                    else:
                        response_content = response_content + '"'
        
        # Fix common LLM formatting issues
        response_content = response_content.replace('\n', ' ').replace('\r', ' ')
        response_content = re.sub(r'\s+', ' ', response_content)
        
        return response_content.strip()
        
    except Exception as e:
        print(f"Warning: JSON cleaning failed: {str(e)}")
        return response_content


def process_chunk_with_gpt4(chunk_text, chunk_index, total_chunks):
    """Process a single chunk of text with GPT-4 using comprehensive extraction rules."""
    prompt = f"""
    You are a survey document parser specialized in extracting complete survey structures. This is chunk {chunk_index + 1} of {total_chunks} from a large survey document.

    IMPORTANT: Extract ALL questions and their content from this chunk, paying special attention to tables, grids, recodes, hidden questions, special question types, and ALL instructions or programming notes.

    FORMATTING PRESERVATION: The text may contain formatting tags like <b>bold</b>, <i>italic</i>, <u>underline</u>, <s>strikethrough</s>, <large>large</large>, <small>small</small>, <br/> (line breaks), and <br/><br/> (paragraph breaks). Always preserve these formatting tags in the extracted question text to maintain the original document styling and layout.

    **CRITICAL: NEVER SKIP REPEATED QUESTIONS - Extract ALL questions in their exact order as they appear in the document, even if they seem identical or repetitive. Each occurrence should be treated as a separate question with its own ID and position.**

    **CRITICAL: NEVER SKIP QUESTIONS WITH IMAGES OR VISUAL CONTENT - Questions containing franchise logos, product images, or visual elements MUST be extracted completely with ALL their options, even if images are represented as placeholders.**

    **DO NOT STOP extraction at any section header, table, or keyword such as 'QUOTAS', 'SCREENER', 'MAIN SURVEY', or similar. Continue extracting ALL questions, instructions, and logic until the very end of the chunk, even if the chunk ends mid-section or mid-table.**
    
    If a section header (e.g., 'QUOTAS', 'MAIN SURVEY', etc.) appears, treat it as a section label and continue extracting all subsequent questions and content. Do not treat any section as the end of the survey. Only stop at the end of the chunk.

    **SECTION HEADER DETECTION:**
    - Look for patterns like "SECTION 1:", "PART A:", "BLOCK 1:", "DEMOGRAPHICS:", etc.
    - Extract section context and hierarchy information
    - Continue processing ALL content after section headers - they are NOT stopping points

    **OPTION LOGIC DETECTION:**
    - **CRITICAL**: Analyze each option for embedded logic, conditions, or special behavior
    - **TERMINATION LOGIC**: Extract "TERMINATE IMMEDIATELY", "TERMINATE IF SELECTED", "DISQUALIFY IF SELECTED"
    - **SKIP LOGIC**: Extract "SKIP TO Q10", "SKIP TO SECTION 2", "GO TO Q15"
    - **CONDITIONAL DISPLAY**: Extract "SHOW Q15 ONLY IF SELECTED", "DISPLAY IF Q1 = 1"
    - **EXCLUSIVE OPTIONS**: Extract "NONE OF THE ABOVE", "OTHER (SPECIFY)", "PREFER NOT TO SAY"
    - **VALIDATION RULES**: Extract "MUST BE 18+", "RANGE 1-99", "REQUIRED IF SELECTED"
    - **DEPENDENCY LOGIC**: Extract "IF SELECTED, SHOW Q10", "ONLY IF Q1 = YES"
    - **RANDOMIZATION**: Extract "RANDOMIZE ORDER", "SHUFFLE OPTIONS"
    - **EXCLUSIVE BEHAVIOR**: Extract options that deselect others when selected
    - **CONDITIONAL TEXT**: Extract piped text that changes based on previous answers
    - **HIDDEN OPTIONS**: Extract options that are only shown under certain conditions
    - **MUTUALLY EXCLUSIVE**: Extract options that cannot be selected together
    - **REQUIRED SELECTIONS**: Extract minimum/maximum selection requirements
    - **CUSTOM VALIDATION**: Extract option-specific validation rules and constraints
    
    **CRITICAL: UNIVERSAL HIDDEN/RECODE VARIABLE DETECTION:**
    - **NEVER SKIP any variable assignment patterns** - these are essential hidden questions for survey logic
    - **UNIVERSAL PATTERN RECOGNITION**: Look for ANY patterns like "VARIABLE_NAME=value IF condition" regardless of prefix
    - **ASSIGNMENT LOGIC EXTRACTION**: Extract complete conditional assignments like "ANY_VAR=1 IF condition"
    - **HIDDEN VARIABLE CREATION**: Convert every assignment pattern to a proper hidden radio question
    - **UNIVERSAL FLAGS**: Most hidden variables are boolean flags (1=Yes, 0=No) or categorical assignments
    - **CONDITIONAL LOGIC**: Extract the full IF-THEN logic regardless of variable naming convention
    - **CROSS-REFERENCES**: Identify variables that reference other questions
    - **MULTI-CONDITION LOGIC**: Handle complex conditions with AND/OR operators universally
    - **QUALIFICATION FLAGS**: Hidden variables often determine survey qualification and flow
    
    **CRITICAL: IMAGE AND VISUAL CONTENT HANDLING:**
    - **NEVER SKIP questions that contain images, logos, or visual elements**
    - **EXTRACT ALL OPTIONS** even if they are in image format or have visual components
    - **IMAGE MARKERS**: Look for [IMAGE], [LOGO], [PICTURE], or similar markers in text
    - **VISUAL SCALES**: Extract questions with images as option choices (logos, product images, etc.)
    - **MIXED CONTENT**: Handle questions with both text and images properly
    - **IMAGE DESCRIPTIONS**: Extract any alt text, captions, or descriptions associated with images
    - **PLACEHOLDER RECOGNITION**: Recognize and process image placeholders in table cells
    
    **CRITICAL: DYNAMIC CONTENT PATTERNS:**
    - **INSERT STUBS SELECTED**: Extract "INSERT STUBS SELECTED IN F1" patterns and convert to proper conditional logic
    - **SHOW IF STUB X SELECTED**: Extract all conditional display logic based on previous selections
    - **SAME ORDER AS**: Extract ordering requirements and maintain sequence relationships
    - **DO NOT SHOW**: Extract exclusion logic for specific options or titles
    - **ROTATE WITH PREVIOUS**: Extract rotation and randomization instructions
    - **CONDITIONAL OPTIONS**: Extract options that only appear based on previous answers

    **COMMENT DETECTION:**
    - **ONLY use comment tags for actual respondent guidance text**
    - **DO NOT use comment tags for programming notes, internal instructions, or technical details**
    - **Examples of what goes in comments:**
      - "Please select all that apply"
      - "Choose the option that best describes you"
      - "You may select more than one answer"
      - "Please be specific in your response"
    - **Examples of what does NOT go in comments:**
      - "PROGRAMMING NOTE: Use checkbox type"
      - "INTERNAL: Track respondent progress"
      - "TECHNICAL: Randomize order"
      - "VALIDATION: Must be 18+"

    **CRITICAL: CDATA SECTION DETECTION:**
    - **DETECT CDATA USAGE**: Look for text containing special characters like "<", "&", ">", or patterns indicating CDATA sections
    - **CDATA INDICATORS**: Text with mathematical operators (<, >, <=, >=), HTML/XML tags, JavaScript code, or special characters
    - **EXTRACT CDATA FLAG**: Set "cdata": true in JSON when question text requires CDATA wrapping
    - **PRESERVE SPECIAL CHARACTERS**: Ensure all special characters are properly extracted and flagged for CDATA wrapping
    
    **CRITICAL: PYTHON VARIABLE SCOPE DETECTION:**
    - **GLOBAL VARIABLES**: Detect variables defined in:
      - extraVariables attribute in survey tag
      - URL parameters (sample source variables)
      - <exec when="init"> blocks (survey-wide constants)
      - <exec when="virtualInit"> blocks (virtual question constants)
    - **LOCAL VARIABLES**: Detect variables in <exec> blocks without when attribute (scoped to current question)
    - **PERSISTENT VARIABLES**: Detect variables with "p." prefix (p.variableName) - per-participant storage
    - **EXTRACT VARIABLE SCOPE**: Add "variable_scope": "global|local|persistent" to questions/variables
    - **EXTRACT EXEC BLOCKS**: Extract all <exec> blocks with:
      - "when" attribute value (init|virtualInit|virtual|none)
      - Complete Python code content
      - Variable assignments and logic
    
    **CRITICAL: CONTROL TAGS DETECTION:**
    - **<goto>**: Extract "SKIP TO", "GO TO", "JUMP TO" patterns → extract target question/section
      - Extract target: question ID or section name
      - Extract condition: when to jump (if conditional)
      - Extract position: where goto appears (after which question)
    - **<exit>**: Extract survey completion actions, finish actions
      - Extract action type: redirect URL, thank you page, etc.
      - Extract condition: when to exit
    - **<finish>**: Extract survey stop/end conditions
      - Extract condition: when survey finishes
      - Extract message: finish message if any
    - **<term>**: Extract termination/disqualification logic (already partially supported)
      - Extract label: term label/identifier
      - Extract condition: when to terminate
      - Extract message: termination message
      - Extract screen: termination screen content
    - **<suspend>**: Extract page breaks (already partially supported)
      - Extract position: after which question
      - Extract condition: conditional page breaks
    - **<block>**: Extract section/block definitions
      - Extract label: block identifier
      - Extract title: block title/name
      - Extract randomize: block-level randomization
      - Extract cond: block conditional display
    - **<condition>**: Extract reusable condition logic references
      - Extract label: condition identifier
      - Extract logic: Python condition expression
      - Extract usage: where condition is referenced
    - **<alert>**: Extract email notification triggers
      - Extract condition: when to send alert
      - Extract recipients: email addresses
      - Extract subject: email subject line
      - Extract message: email content
    - **<loop>/<looprow>**: Extract looping/cycling patterns
      - Extract label: loop identifier
      - Extract rows: rows to loop through
      - Extract questions: questions inside loop
      - Extract randomize: loop randomization
    - **<insert>**: Extract grid splitting patterns
      - Extract target: target question ID
      - Extract source: source question ID
      - Extract condition: when to insert
    - **<pipe>**: Extract conditional information display
      - Extract condition: when to show piped content
      - Extract content: piped text/variable
      - Extract format: formatting instructions
    - **EXTRACT ALL**: Add "control_tags" array to survey-level JSON with all detected control tags
    - **POSITION TRACKING**: Track where each control tag appears (after which question)

    CRITICAL QUESTION TYPE CLASSIFICATION:
    For each question, analyze its content and intent to determine the most appropriate type. Use ONLY these exact Decipher XML types:

    | Question Type/Intent                | Decipher XML Type (type value) |
    |-------------------------------------|:------------------------------:|
    | Single Choice                       | radio                          |
    | Multiple Choice                     | checkbox                       |
    | Checkbox                            | checkbox                       |
    | Dichotomous                         | radio                          |
    | Rating Scale                        | radio                          |
    | Likert Scale                        | radio                          |
    | Matrix/Grid                         | grid                           |
    | Dropdown                            | select                         |
    | Slider                              | number                         |
    | Open-Ended Text (single line)       | text                           |
    | Open-Ended Text (multi-line)        | textarea                       |
    | Floating Point Number               | float                          |
    | Virtual/Calculated                  | virtual                        |
    | Logic Node                          | logic                          |
    | Autosuggest/Auto-complete           | autosuggest                    |
    | Button Select                       | button_select                  |
    | Button Rating                       | button_rating                  |
    | Card Rating                         | card_rating                    |
    | Ranking                             | radio                          |
    | Demographic                         | radio                          |
    | Visual Scale                        | radio                          |
    | Benchmarkable                       | radio                          |
    | Small Text Box                      | text                           |
    | ANY Numeric Input (including):      | number                         |
    | - Age, years, months, weeks, days   | number                         |
    | - Quantities, amounts, counts       | number                         |
    | - Percentages, rates, scores        | number                         |
    | - Measurements, values, ranges      | number                         |
    | - Currency amounts, prices          | number                         |
    | - Dates, times, durations           | number                         |
    | - Any question requiring a number   | number                         |
    | ANY Decimal/Float Input:            | float                          |
    | - Decimal percentages               | float                          |
    | - Precise measurements              | float                          |
    | - Currency with decimals            | float                          |
    | - Any question requiring decimals   | float                          |

    CRITICAL RULES FOR TYPE CLASSIFICATION:
    1. For ANY question that requires a numeric response, ALWAYS use 'number' type
    2. Consider both explicit and implicit numeric requirements
    3. Look for context clues about the expected response type
    4. Consider common survey patterns and conventions
    5. When in doubt about numeric input, prefer 'number' type
    6. For questions that could be either numeric or text, prefer 'number' if numeric is more appropriate

    CRITICAL EXTRACTION RULES:
    0. **NEVER USE GENERIC EXAMPLES - ALWAYS EXTRACT ACTUAL LOGIC FROM QUESTION TEXT**
       - Do NOT copy example logic like "IF P1.r1 AND S3A.any AND S5.r1 THEN QUALIFIED"
       - Do NOT use placeholder text like "Add logic here" or "Example logic"
       - ALWAYS read the actual question text and extract the real conditions, logic, and rules
       - If the question says "To qualify as US Homeowner, respondents must: • Preload (P1/1) • Consent to survey (C1/1) • 18+ (S1A/2-7) • Live in the US (S3a_Recode/1-4) • Not working in Advertising/Marketing/Media (S7 ≠ 1) • Be a homeowner (S12/1) • Have a homeowner's insurance policy (S15/2)" then extract EXACTLY that logic
    
    **CRITICAL MISSING QUESTION DETECTION:**
    - **P1. PRELOADS**: Extract as preload question with ID "P1" and ALL options
    - **CONSENT C1**: Extract consent questions starting with "C1" or in "CONSENT" sections  
    - **SCREENER SINT**: Extract screener intro questions starting with "SINT" or in "SCREENER" sections
    - **CASE INSENSITIVE**: "PRELOADS" = "preloads", "CONSENT" = "consent", "SCREENER" = "screener"
    
    1. Extract ALL question types including:
       - Standard questions (single choice, multiple choice, open-ended)
       - Checkbox questions
       - Dichotomous questions (Yes/No, True/False)
       - Grid/Matrix questions
       - Scale questions (rating, likert, semantic, etc.)
       - Dropdown questions
       - Slider questions
       - Ranking questions
       - Demographic questions (age, gender, etc.)
       - Visual scale questions (images, icons, etc.)
       - Benchmarkable questions
       - Small text box questions (single-line) - use type="text"
       - Multi-line text questions - use type="textarea"
       - Numeric input questions (integers) - use type="number"
       - Floating point/decimal questions - use type="float"
       - Virtual/Calculated questions - use type="virtual" with calculation logic
       - Logic node questions - use type="logic"
       - Autosuggest questions - use type="autosuggest"
       - Button Select questions - use type="button_select"
       - Button Rating questions - use type="button_rating"
       - Card Rating questions - use type="card_rating"
       - Hidden questions (often marked as "HIDDEN" or "INTERNAL")
       - Recode questions (often marked as "RECODE" or "COMPUTED")
       - Preload questions (pre-populated or prefilled variables)
       - Derived or computed variables
       - Piped questions (questions that reference previous answers)
       - Logic blocks and skip patterns
       - Quota questions
       - Screening questions
       - Validation questions
       - Randomization blocks
       - Rotation blocks
       - Custom variables and computed fields
       - Note elements (programmer comments) - extract as metadata
    17. **FOR QUOTA QUESTIONS: Apply INTELLIGENT quota pattern recognition** 
        - **ANALYZE the quota pattern first** before creating questions:
          
        **PATTERN 1 - COMPETING SEGMENTS**: Multiple audience qualifications (S100A, S100B, S100C)
        → COMBINE into single question (S100) with if/elif/else logic using <radio> type
        → Each audience becomes a row option, respondent qualifies for exactly ONE
        
        **PATTERN 2 - INDEPENDENT FLAGS**: Individual qualification criteria that can coexist  
        → Keep as separate <checkbox> questions if they represent independent boolean flags
        
        **PATTERN 3 - DEMOGRAPHIC COPYING**: Quotas mirroring demographic breakdowns
        → Create single <radio> question with simple assignment: QuotaID.val = SourceID.val
        
        **PATTERN 4 - CONDITIONAL DEMOGRAPHICS**: Audience-specific demographic splits (S101a=Male Homeowners, S101b=Male Protectors)
        → Create single <radio> with conditional logic based on audience + demographic
        
        - **INTELLIGENT COMBINATION LOGIC**: Look at question IDs and content to determine combination strategy
        - **QUOTA TYPE SELECTION**: Use <radio> for exclusive choices, <checkbox> for independent flags
        - Do NOT use generic examples - extract the REAL logic and apply appropriate pattern
        - **QUOTA CONDITION LOGIC CONVERSION RULES**:
          - Convert question references (QID/value) to Python syntax (QID.rValue)
          - Convert ranges (QID/1-4) to OR logic (QID.r1 or QID.r2 or QID.r3 or QID.r4)
          - Convert negations (QID ≠ value) to not logic (not QID.rValue)
          - Join multiple conditions with 'and' operator
          - Extract ALL qualification conditions into quota_info.qualification_logic
        - Always include: <suspend/>, <term> (if termination needed), <quota> tags after each quota question
    
    18. **FOR PRELOAD QUESTIONS: Extract the ACTUAL source and validation from the question text**
        - If the question says "Set from URL parameter 'preload'" then extract: "source": "URL parameter 'preload'"
        - If the question says "Must be 1 or 2, terminate if missing" then extract: "validation": "Must be 1 or 2, terminate if missing"
        - Do NOT use generic examples - extract the REAL requirements from the question
    
    19. DO NOT MISS ANY QUESTION OR QUESTION-RELATED CONTENT. This includes:
        - All visible, hidden, derived, computed, preload, and recode questions (such as RH, preload variables, etc.)
        - Any variable, question, or item referenced in logic, quotas, instructions, programming notes, or as a variable name
        - Questions or variables that are only shown conditionally, or only referenced in logic, recoding, or programming
        - Any content, label, or variable that could be used for data collection, processing, or analysis
        - **CRITICAL: ALL HIDDEN VARIABLE PATTERNS** - Extract every variable assignment pattern as a separate hidden question
        - If in doubt, ALWAYS include the item as a question or variable and provide as much detail as possible
        - Err on the side of inclusion: treat anything that could possibly be a question, variable, or data point as such, regardless of its format, location, or presentation in the document
        
    **CRITICAL: UNIVERSAL HIDDEN VARIABLE EXTRACTION FORMAT:**
    For ANY variable assignment pattern found in the text (VAR_NAME=value IF condition), extract it as a hidden question:
    ```json
    {{
        "id": "VARIABLE_NAME",
        "text": "HIDDEN QUESTION - VARIABLE_NAME=value IF condition",
        "type": "hidden",
        "purpose": "derived_variable",
        "recode_info": {{
            "source_questions": ["source_question"],
            "logic": "IF condition",
            "assignment": "VARIABLE_NAME=value",
            "condition_type": "boolean_flag"
        }},
        "options": [
            {{"value": 1, "text": "Condition met"}},
            {{"value": 0, "text": "Condition not met"}}
        ]
    }}
    ```

    2. For recode questions, capture:
       - Original question reference  
       - Recode logic and conditions
       - Value mappings
       - Target variable names
       - **CRITICAL: Generate ACTUAL Python logic in exec blocks for mapping values**
       - **NEVER use placeholder comments like <!-- Logic here -->**
       - **For regional coding: Map states/provinces to their correct regional categories**

    3. For hidden questions, capture:
       - Question ID and text
       - Purpose (e.g., internal tracking, validation)
       - Associated visible questions
       - Any special processing rules

    CONDITIONS & INSTRUCTIONS EXTRACTION RULE:
    - For every question, if you see any text that describes when or to whom the question should be shown (e.g., "ASK IF...", "SHOW IF...", "ONLY IF...", "BLOCK IF...", "TERMINATE IF..."), extract this as a 'conditions' field.
    - If you see any text that gives instructions to the respondent, interviewer, or programmer (e.g., "Please select all that apply", "Disqualify if...", "Randomize order", "Do not read out", "Instruction:", "Note:"), extract this as an 'instructions' field.
    - Use your best judgment and natural language understanding to distinguish between logic/conditions and instructions, and do not omit any such content.
    - If in doubt, err on the side of inclusion: better to include a note in 'instructions' or 'conditions' than to miss it.
    - Examples:
      - "ASK IF COUNTRY = US" → "conditions": "ASK IF COUNTRY = US"
      - "Please select all that apply" → "instructions": ["Please select all that apply"]
      - "Disqualify if selected 'Prefer not to say'" → "instructions": ["Disqualify if selected 'Prefer not to say'"]
      - "SHOW IF Q1 = 1" → "conditions": "SHOW IF Q1 = 1"
    
### Instruction Processing
- **Termination instructions** (e.g., "Disqualify if...", "Terminate if..."): Convert to `<term>` tags
- **Validation instructions** (e.g., "Flag if...", "Check if..."): Convert to hidden validation questions with `<exec>` logic
- **Display instructions** (e.g., "Show only if...", "Ask if..."): Convert to `cond` attributes
- **Randomization instructions** (e.g., "Randomize options", "Shuffle rows"): Add appropriate shuffle attributes

    4. For grid/matrix questions, format as:
       {{
           "id": "Q1",
           "text": "Question stem",
           "type": "grid",
           "rows": ["Row item 1", "Row item 2", ...],
           "columns": ["Column 1", "Column 2", ...],
           "table_data": {{
               "headers": [...],
               "rows": [...]
           }},
           "rotation": "none|rows|columns|both",
           "randomization": "none|rows|columns|both"
       }}

    5. For scale questions, include all scale points:
       {{
           "id": "Q2",
           "text": "Question text",
           "type": "radio",
           "scale_points": [
               {{"value": "1", "label": "Strongly Disagree"}},
               {{"value": "2", "label": "Disagree"}},
               ...
           ],
           "scale_type": "likert|semantic|numeric|custom"
       }}

    6. For recode questions, format as:
       {{
           "id": "RECODE1",
           "text": "Recode description",
           "type": "radio",
           "source_question": "Q1",
           "logic": "IF Q1 = 1 THEN 1 ELSE 0",
           "value_mappings": [
               {{"source": "1", "target": "1"}},
               {{"source": "2", "target": "0"}}
           ],
           "target_variable": "Q1_RECODE"
       }}

    7. For hidden questions, format as:
       {{
           "id": "HIDDEN1",
           "text": "Hidden question text",
           "type": "radio",
           "purpose": "validation|tracking|internal",
           "associated_questions": ["Q1", "Q2"],
           "processing_rules": "Special instructions if any"
       }}

    8. For virtual questions, format as:
       {{
           "id": "VIRTUAL1",
           "text": "Virtual question description",
           "type": "virtual",
           "virtual": "calculation_logic_here",
           "source_questions": ["Q1", "Q2"],
           "calculation": "Python expression or formula",
           "where": "report|survey|execute"
       }}

    9. For questions with exec blocks, format as:
       {{
           "id": "Q1",
           "text": "Question text",
           "type": "radio",
           "exec_blocks": [
               {{
                   "when": "init|virtualInit|virtual|none",
                   "code": "Python code here",
                   "purpose": "variable_assignment|calculation|logic"
               }}
           ],
           "variable_scope": "global|local|persistent"
       }}

    10. For questions with validation, format as:
        {{
            "id": "Q1",
            "text": "Question text",
            "type": "number",
            "validation": {{
                "rules": [
                    {{"type": "range", "min": 1, "max": 100}},
                    {{"type": "required", "message": "This field is required"}}
                ],
                "messages": ["Error message 1", "Error message 2"],
                "conditions": "IF condition THEN validate"
            }}
        }}

    11. For questions with nets (aggregation), format as:
        {{
            "id": "Q1",
            "text": "Question text",
            "type": "checkbox",
            "nets": [
                {{
                    "label": "Net1",
                    "title": "Net Title",
                    "type": "sum|average|count|percent",
                    "rows": ["r1", "r2", "r3"],
                    "columns": ["c1", "c2"],
                    "choices": ["ch1", "ch2"],
                    "condition": "conditional net display",
                    "where": "report|survey"
                }}
            ]
        }}
    
    12. For questions with vector logic, format as:
        {{
            "id": "Q1",
            "text": "Question text",
            "type": "checkbox",
            "conditions": "Q1.any and Q2.r1",
            "vector_logic": true,
            "vector_operations": [
                {{"operation": "any", "question": "Q1"}},
                {{"operation": "count", "question": "Q1", "comparison": "gt", "value": 3}}
            ]
        }}
    
    13. For questions with advanced randomization, format as:
        {{
            "id": "Q1",
            "text": "Question text",
            "type": "grid",
            "randomization": {{
                "type": "rows|columns|both|balanced|latin_square|block",
                "shuffle": "rows|columns|both",
                "condition": "conditional randomization",
                "rotation": "rotate_with_previous|same_order_as|insert_downs",
                "pattern": "randomization pattern description"
            }}
        }}

    14. For questions with complete attributes, format as:
        {{
            "id": "Q1",
            "text": "Question text",
            "type": "radio",
            "cdata": false,
            "where": "survey|report|execute",
            "attributes": {{
                "optional": "0|1",
                "unique": "0|1",
                "verify": "verification_rule",
                "size": "input_size",
                "atleast": "minimum_selections",
                "atmost": "maximum_selections",
                "shuffle": "rows|columns|both|none",
                "randomize": "randomization_pattern",
                "cond": "conditional_display_logic",
                "di": "display_instruction",
                "ss": "single_select_mode",
                "adim": "report_dimension",
                "virtual": "virtual_calculation"
            }},
            "cell_attributes": {{
                "rows": [
                    {{
                        "label": "r1",
                        "open": "0|1",
                        "exclusive": "0|1",
                        "cond": "conditional_display",
                        "value": "explicit_value"
                    }}
                ],
                "columns": [...],
                "choices": [...]
            }}
        }}

    13. For EVERY question, extract ALL instructions, programming notes, interviewer notes, or special directives (such as "RANK UP TO 5", "AUTOCODE AS RANK 1", "INSERT DOWNS", etc.) and place them in an 'instructions' field. Do not miss any such content, even if it is outside the main question text or in a different color/font/style.

    14. For survey-level elements, extract:
        - XML declaration (version, encoding)
        - Survey attributes (extraVariables, compat, theme, style, etc.)
        - Control tags (goto, exit, finish, term, suspend, block, condition, alert, loop, insert, pipe)
        - Python variables by scope (global, local, persistent)
        - Styling information (style, stylevar, theme, themevars, CSS/JS files)
        - Participant sources (samplesource, samplesources, URL variables)
        - Note elements (programmer comments)
    
    **CRITICAL: STYLING SYSTEM EXTRACTION:**
    - **<style> TAG**: Extract all `<style>` tag content (CSS code)
      - Extract complete CSS content
      - Extract style name/identifier if present
      - Extract where attribute (survey|report|execute)
      - Extract any conditions for style application
    - **<stylevar> TAG**: Extract all `<stylevar>` variables
      - Extract variable name
      - Extract variable value
      - Extract variable type/usage
    - **<theme> TAG**: Extract all `<theme>` tag content
      - Extract theme name/identifier
      - Extract theme content (CSS/Less code)
      - Extract theme attributes
    - **<themevars> TAG**: Extract all `<themevars>` variables
      - Extract variable name
      - Extract variable value
      - Extract variable scope
    - **CSS FILES**: Detect CSS file references
      - Look for .css file paths
      - Look for CSS file imports
      - Extract file paths and locations
    - **JAVASCRIPT FILES**: Detect JavaScript file references
      - Look for .js file paths
      - Look for JavaScript file imports
      - Extract file paths and locations
    - **STATIC FILES**: Detect static file references (images, fonts, etc.)
      - Look for file references in styling context
      - Extract file paths and types
    - **LESS CSS SYSTEM**: Detect Less CSS usage
      - Look for Less CSS syntax (@variables, nesting, mixins)
      - Extract Less CSS content
      - Extract Less variables
    - **GOOGLE FONTS**: Detect Google Fonts imports
      - Look for fonts.googleapis.com references
      - Extract font family names
      - Extract font import URLs
    
    **CRITICAL: PARTICIPANT SOURCE CONFIGURATION:**
    - **<samplesource> TAG**: Extract samplesource configurations
      - Extract samplesource name/identifier
      - Extract source type (URL, panel, API, etc.)
      - Extract source configuration details
      - Extract validation rules
      - Extract variable mappings
    - **<samplesources> TAG**: Extract multiple source configurations
      - Extract all samplesource definitions
      - Extract source priorities/order
      - Extract source conditions
    - **<var> TAG**: Extract URL variable definitions
      - Extract variable name
      - Extract variable values/options
      - Extract variable validation rules
      - Extract variable default values
    - **PARTICIPANT LOCKING**: Extract locking mechanisms
      - Detect unique variable locking patterns
      - Extract locking variable names
      - Extract locking conditions
      - Extract locking messages
    - **ONE-TIME ACCESS**: Extract access restrictions
      - Detect "take survey only once" patterns
      - Extract access control rules
      - Extract restriction conditions
    - **COMPLETION ACTIONS**: Extract finish/exit actions
      - Extract completion redirect URLs
      - Extract completion messages
      - Extract completion conditions
    
    **CRITICAL: ENHANCED SURVEY ATTRIBUTES EXTRACTION:**
    - **PROJECT SETTINGS**: Extract all project-level settings
      - extraVariables: comma-separated variable list
      - compat: compatibility level (e.g., compat="28")
      - survey mode: live, test, preview modes
      - language settings: default language, multi-language configs
      - project-level configurations
    - **DISPLAY SETTINGS**: Extract all display-related settings
      - Theme settings: theme name, theme configuration
      - Style settings: style name, style configuration
      - Mobile optimization: mobile-specific display settings
      - Browser settings: browser-specific configurations
      - Display mode settings
    - **DEVICE SETTINGS**: Extract device-specific configurations
      - Mobile device settings
      - Tablet device settings
      - Desktop device settings
      - Device detection settings
    - **FIELD SETTINGS**: Extract field-level defaults
      - Default field sizes
      - Default field validation
      - Default field styling
      - Field-level configurations
    - **QUESTION SETTINGS**: Extract question-level defaults
      - Default question attributes
      - Default question styling
      - Default question behavior
      - Question-level configurations
    - **SYSTEM LANGUAGE RESOURCES**: Extract system text
      - Default error messages
      - Default tooltips
      - Default button text
      - Browser window titles
      - Support link text
      - System language overrides
    
    **CRITICAL: DATASOURCE TAG EXTRACTION:**
    - **<datasource> TAG**: Extract datasource configurations
      - Extract file path (tab-delimited file)
      - Extract column mappings
      - Extract data source name
      - Extract data source type
      - Extract refresh settings
      - Extract data validation rules
    
    **CRITICAL: DEFINE TAG EXTRACTION:**
    - **<define> TAG**: Extract reusable answer list definitions
      - Extract define name/identifier
      - Extract answer list content (rows, columns, choices)
      - Extract define attributes
      - Extract usage locations (where define is referenced)
    
    **CRITICAL: HTML TAG EXTRACTION:**
    - **<html> TAG**: Extract HTML comment elements
      - Extract HTML content
      - Extract where attribute (survey|report|execute)
      - Extract HTML conditions
      - Extract HTML formatting
    
    **CRITICAL: NOTE TAG EXTRACTION:**
    - **<note> TAG**: Extract programmer comment elements
      - Extract note content
      - Extract note placement (where it appears)
      - Extract note conditions
      - Extract note formatting
    
    **CRITICAL: MACRO SYSTEM DETECTION:**
    - **MACRO DEFINITIONS**: Detect `<macro>` tag definitions
      - Extract macro name/identifier
      - Extract macro parameters
      - Extract macro content/template
      - Extract macro attributes
    - **MACRO USAGE**: Detect macro invocations
      - Extract macro call locations
      - Extract macro parameters passed
      - Extract macro expansion context
    - **MACRO EXPANSION**: Extract macro expansion logic
      - Extract how macros are expanded
      - Extract macro variable substitutions
    
    **CRITICAL: MUTATOR DETECTION:**
    - **MUTATOR FUNCTIONS**: Detect mutator function definitions
      - Extract mutator function names
      - Extract mutator parameters
      - Extract mutator logic
    - **SURVEY STRUCTURE MODIFICATIONS**: Detect structure changes
      - Extract question creation patterns
      - Extract question modification patterns
      - Extract row/column creation patterns
      - Extract structure transformation logic
    
    **CRITICAL: DATABASE INTEGRATION DETECTION:**
    - **ADB (AUTOMATED DATABASE SYSTEM)**: Detect ADB usage
      - Extract ADB configuration
      - Extract ADB data sources
      - Extract ADB integration patterns
    - **AUXILIARY DATABASE**: Detect auxiliary database usage
      - Extract auxiliary database configuration
      - Extract database read/write operations
      - Extract database connection details
    - **loadData/loadRecord**: Detect data loading functions
      - Extract loadData() function calls
      - Extract loadRecord() function calls
      - Extract data source references
      - Extract data loading conditions
    
    **CRITICAL: TRANSFORMATION SYSTEM DETECTION:**
    - **TRANSFORMATION SCRIPTS**: Detect transformation script definitions
      - Extract script name/identifier
      - Extract script content
      - Extract script triggers
    - **TRANSFORMATION TRIGGERS**: Detect when transformations run
      - Extract completion-based triggers
      - Extract event-based triggers
      - Extract condition-based triggers
    - **TRANSFORMATION LOGIC**: Extract transformation operations
      - Extract content modifications
      - Extract structure changes
      - Extract data transformations
    
    **CRITICAL: HIDDEN QUESTION DETECTION:**
    Use basic pattern matching to detect hidden questions:
    
    **BASIC INDICATORS:**
    - Questions about "qualification", "eligibility", "derived values", "computed results"
    - Text containing logical operations: "IF-THEN", "based on", "calculated using", "determined by"
    - Administrative purposes: "tracking", "internal", "system", "flag", "counter"
    - Assignment language: "set to", "equals", "assigned", "computed as"
    
    **QUESTION BOUNDARY DETECTION:**
    - Look for natural language cues: question numbers, section breaks, topic shifts
    - Detect transition words: "next", "now", "moving on", "question X"
    - Identify format changes: numbering systems, indentation, styling
    - Recognize logical groupings: related content vs distinct questions
    
    **VARIABLE DETECTION:**
    - ANY text describing variable assignment (regardless of naming convention)
    - Computational logic embedded in instructions or notes
    - Conditional statements that set values based on responses
    - Rules that determine qualification, segmentation, or categorization

    QUOTA AND PRELOAD QUESTION EXTRACTION RULES:
    - **Quota Questions**: Look for questions with IDs like S100, S100A, S100B, or text containing "quota", "qualification", "overall qualifications", "screening", etc. These are critical for survey flow control.
    - **Preload Variables**: Look for questions with IDs like P1, P1_PRELOAD, or text containing "preload", "pre-populated", "prefilled", "URL variable", etc. These are passed via URL parameters.
    - **Hidden Questions**: Look for questions marked as "HIDDEN", "INTERNAL", "NOT SHOWN TO RESPONDENT", or similar indicators.
    - **Recode Questions**: Look for questions marked as "RECODE", "COMPUTED", "DERIVED", or similar indicators.
    - **Termination Logic**: Look for any text indicating when respondents should be terminated, disqualified, or screened out.
    
    **STRUCTURE EXTRACTION:**
    - **Conditions**: Look for [ASK IF X], [SHOW IF X], [SKIP TO X] patterns and extract display/skip logic
    - **Instructions**: Look for [MULTI SELECT], [SINGLE SELECT], [EXCLUSIVE], [TERMINATE IF X] programming notes
    - **Comments**: Look for "Please select all that apply", "Choose one", etc. respondent guidance text
    - **Options**: Extract option-specific logic like "Option 4: [TERMINATE IF P1/3]"
    - **Nested Logic**: Capture complex IF-THEN-ELSE structures and question dependencies
    
    **CRITICAL: VECTOR LOGIC DETECTION:**
    - **VECTOR SYNTAX**: Detect abbreviated Python condition logic using vector operations
    - **PATTERNS TO DETECT**:
      - Q1.any (any row selected in Q1)
      - Q1.all (all rows selected in Q1)
      - Q1.none (no rows selected in Q1)
      - Q1.count (count of selected rows)
      - Q1.sum (sum of selected values)
      - Q1.rows (all rows in question)
      - Q1.cols (all columns in question)
    - **VECTOR OPERATIONS**: Extract vector-based conditions like "Q1.any and Q2.r1"
    - **CONVERT TO STANDARD**: When extracting, preserve vector syntax but also note it's vector logic
    - **EXAMPLES**:
      - "IF Q1.any THEN show Q2" → Extract as vector condition
      - "IF Q1.count gt 3 THEN terminate" → Extract vector count operation
      - "IF Q1.sum lt 10 THEN show Q3" → Extract vector sum operation
    
    **CRITICAL: ENHANCED RANDOMIZATION PATTERNS:**
    - **ROW RANDOMIZATION**: Extract "RANDOMIZE ROWS", "SHUFFLE ROWS", "ROTATE ROWS"
    - **COLUMN RANDOMIZATION**: Extract "RANDOMIZE COLUMNS", "SHUFFLE COLUMNS", "ROTATE COLUMNS"
    - **CHOICE RANDOMIZATION**: Extract "RANDOMIZE CHOICES", "SHUFFLE CHOICES"
    - **CONDITIONAL RANDOMIZATION**: Extract "RANDOMIZE IF condition", "SHUFFLE WHEN X"
    - **ROTATION PATTERNS**: Extract "ROTATE WITH PREVIOUS", "SAME ORDER AS", "INSERT DOWNS"
    - **RANDOMIZATION ATTRIBUTES**: Extract shuffle="rows|columns|both", randomize="pattern"
    - **ADVANCED PATTERNS**: Extract balanced rotation, Latin square, block randomization
    
    **CRITICAL: ENHANCED IMAGE METADATA EXTRACTION:**
    - **IMAGE SOURCES**: Extract image file paths, URLs, or references
    - **IMAGE ATTRIBUTES**: Extract alt text, title, description, caption
    - **IMAGE SIZING**: Extract width, height, size constraints
    - **IMAGE POSITIONING**: Extract alignment, positioning (left, right, center, inline)
    - **IMAGE MARKERS**: Detect [IMAGE], [LOGO], [PICTURE], [PHOTO] placeholders
    - **VISUAL CONTENT**: Extract questions with images as options (product images, logos, etc.)
    - **MIXED CONTENT**: Handle questions with both text and images properly
    - **IMAGE DESCRIPTIONS**: Extract any alt text, captions, or descriptions associated with images
    - **PLACEHOLDER RECOGNITION**: Recognize and process image placeholders in table cells
    - **IMAGE FORMATS**: Detect .jpg, .png, .gif, .svg, .webp references
    
    **CRITICAL: COMPLETE ATTRIBUTE EXTRACTION PATTERNS:**
    - **ELEMENT ATTRIBUTES** (all questions):
      - where: "survey|report|execute" - where element appears
      - di: display instruction text
      - cond: conditional display logic (Python expression)
      - randomize: randomization pattern
      - shuffle: "rows|columns|both|none"
      - optional: "0|1" - required vs optional
      - unique: "0|1" - unique value requirement
      - verify: verification rule
      - size: input field size
      - ss: single-select mode
      - browser: browser-specific settings
      - compat: compatibility level
      - mobile: mobile-specific settings
    - **QUESTION ATTRIBUTES** (question-specific):
      - adim: primary report dimension
      - virtual: virtual calculation logic
      - markers: marker assignment
      - atleast: minimum selections required
      - atmost: maximum selections allowed
      - open: open-ended option flag
      - exclusive: mutually exclusive option flag
    - **CELL ATTRIBUTES** (rows/columns/choices):
      - open: "0|1" - open-ended cell option
      - exclusive: "0|1" - mutually exclusive option
      - randomize: cell-level randomization
      - cond: cell-level conditional display
      - value: explicit value assignment
      - label: cell label/identifier
    - **EXTRACT ALL**: Ensure every attribute mentioned in instructions is extracted
               
    CRITICAL: Do NOT skip any quota, preload, hidden, or recode questions. These are essential for survey functionality even if they don't appear to respondents directly.
    
    QUOTA DETECTION:
    - Look for questions with IDs starting with "S100" (S100, S100A, S100B, S100C, etc.)
    - Look for questions containing keywords: "quota", "qualification", "screening", "overall qualifications", "respondent qualification", "survey qualification", "participant qualification"
    - Look for questions that determine survey eligibility or continuation
    - Look for questions that reference "N=" targets or termination conditions
    - Look for questions that control survey flow based on previous responses
    
    UNIVERSAL QUOTA LOGIC PRINCIPLES:
    
    **QUOTA QUESTION TYPE DETECTION:**
    1. **AUDIENCE/SEGMENT QUALIFICATION (use <radio>):** Multiple mutually exclusive audience segments where respondent can qualify for only ONE
       - Pattern: Questions like "S100A To qualify as [Audience A]", "S100B To qualify as [Audience B]"
       - Logic: Use if/elif/else structure - respondent gets assigned to exactly one segment or disqualified
       - Structure: Combine all segments into ONE question with multiple rows for each audience + "Not qualified" row
    
    2. **INDIVIDUAL QUALIFICATION FLAGS (use <checkbox>):** When each qualification is independent and multiple can be true
       - Pattern: Questions that check individual criteria that can coexist
       - Logic: Each checkbox represents an independent boolean flag
       - Structure: Keep as separate questions if they represent truly independent qualifications
    
    3. **DEMOGRAPHIC COPYING (use <radio>):** Quotas that copy values from previously answered questions
       - Pattern: Questions like "S101 Gender Quotas", "S102 Age Quotas"  
       - Logic: Simple assignment from source question (S101.val = S2.val)
       - Structure: Single question that mirrors the source question structure
    
    **UNIVERSAL QUOTA COMBINATION RULES:**
    - **Multiple segments competing for one slot:** Combine into single <radio> with if/elif/else
    - **Independent qualifications:** Keep as separate <checkbox> questions
    - **Demographic breakdowns:** Create single <radio> that copies from source question
    - **Complex conditional assignments:** Use <radio> with conditional logic in <exec>
    
    **QUOTA QUESTION ID PATTERN ANALYSIS:**
    - IDs ending in letters (S100A, S100B, S100C) → Usually competing segments → Combine into base ID (S100)
    - IDs with same prefix but different suffixes (S101a, S101b, S101c) → Usually demographic splits → Combine into base ID (S101)
    - Sequential numbered IDs (S100, S101, S102) → Usually separate quota dimensions → Keep separate
    - Complex suffixes (S104c, S105ab) → Usually conditional/filtered quotas → Single question with conditional logic
    
    **AUTOMATIC QUOTA STRUCTURE GENERATION:**
    1. **Identify quota type** based on question content and ID patterns
    2. **Determine combination strategy** (combine competing segments vs keep independent)
    3. **Generate appropriate tag type** (<radio> for exclusive, <checkbox> for independent)
    4. **Create logic structure** (if/elif/else for segments, simple assignment for demographics)
    5. **Add required components** (<suspend/>, <term> if needed, <quota> tags)
    
    **EXEC BLOCK LOGIC GENERATION RULES:**
    - **NEVER use placeholder comments or <!-- --> in exec blocks**
    - **ALWAYS generate functional Python code that executes correctly**
    - **USE ONLY JSON CONTENT - no generic/placeholder logic**
    - **Quota questions:** Use actual qualification_logic from quota_info in JSON
    - **Demographic copying:** Simple assignment: ID.val = SourceID.val  
    - **Recode questions:** Use ACTUAL recode_info.logic from JSON, convert to proper Python syntax
    - **Recode conversion examples:**
      - JSON: "IF Q1=1 THEN 1 ELSE 2" → Python: "if Q1.r1: ID.val = ID.r1.index else: ID.val = ID.r2.index"
      - JSON: "Map states to regions" → Use actual state mapping from JSON, not generic
    - **Regional coding:** Use actual state-to-region mappings from JSON recode_info
    - **Boolean flags:** Use ID.rX.val = 1 or ID.rX.val = 0 assignments based on JSON logic
    
    PRELOAD DETECTION:
    - Look for questions with IDs starting with "P" (P1, P1_PRELOAD, P2, etc.)
    - Look for sections titled "PRELOADS", "P1. PRELOADS", "URL PARAMETERS", "INVITATION VARIABLES"
    - Look for questions containing keywords: "preload", "preloads", "pre-populated", "prefilled", "URL variable", "parameter", "invitation code", "sample source"
    - Look for hidden questions that get values from external sources
    - Look for questions that reference URL parameters or invitation links
    - Look for questions marked as "HIDDEN" or "INTERNAL" that set respondent attributes
    - **CRITICAL**: Case-insensitive detection - "PRELOADS" = "preloads" = "Preloads"
    
    HIDDEN QUESTION DETECTION:
    - Look for questions marked as "HIDDEN", "INTERNAL", "NOT SHOWN TO RESPONDENT", "PROGRAMMER ONLY"
    - Look for questions that are used for logic, validation, or tracking purposes
    - Look for questions that compute values based on other responses
    - Look for questions that control survey flow without being visible to respondents
    
    RECODE DETECTION:
    - Look for questions marked as "RECODE", "COMPUTED", "DERIVED", "CALCULATED"
    - Look for questions that create new variables based on existing responses
    - Look for questions that combine or transform previous answers
    - Look for questions that create summary or aggregate variables
    
    QUOTA SHEET AND PRELOAD DEFINITION EXTRACTION:
    - Look for sections titled "QUOTAS", "QUOTA SHEET", "SAMPLE TARGETS", "RESPONDENT TARGETS"
    - Look for tables or lists showing quota targets (e.g., "N=1,600", "Target: 800", "Limit: 200")
    - Look for sections titled "PRELOADS", "P1. PRELOADS", "URL PARAMETERS", "INVITATION VARIABLES", "SAMPLE SOURCE VARIABLES"
    - Look for variable definitions like `<var name="preload" values="1,2,3,4,5">`
    - Look for any text describing how preload variables are used or validated
    - Extract all quota targets, limits, and qualification criteria
    - Extract all preload variable definitions, valid values, and validation rules
    - **CONSENT SECTION DETECTION**: Look for "CONSENT", "C1.", consent text, data protection notices
    - **SCREENER SECTION DETECTION**: Look for "SCREENER", "SINT.", "Thank you for participating", screener intro text
    - **NUMERIC QUESTION DETECTION**: Look for age questions, quantity questions, percentage questions - always use type="number"

    SPECIAL MULTI-COUNTRY EXTRACTION RULE:
    - If a question contains multiple country-specific blocks (e.g., 'Ask US only', 'Ask UK only', 'Ask DE only', etc.), you MUST extract each block as a SEPARATE question object in the output JSON. Each should have its own id, text, options (with full lists if referenced), conditions, and instructions. Do NOT combine them into a single question or as subquestions—treat each as a distinct question for its country. This ensures that all country-specific variants are fully represented in the extracted JSON.

    **QUOTA LOGIC EXTRACTION RULES:**
    For quota questions, convert qualification text to Python logic using these transformations:
    - (QID/value) → QID.rValue
    - (QID/value1-valueN) → (QID.rValue1 or QID.rValue2 or ... or QID.rValueN)  
    - (QID ≠ value) → not QID.rValue
    - Multiple conditions → join with 'and'
    - Extract into quota_info.qualification_logic field
    
    Output format should be a JSON object with survey-level metadata and questions array:
    {{
        "survey": {{
            "xml_declaration": {{
                "version": "1.0",
                "encoding": "UTF-8|ISO-8859-1|UTF-16"
            }},
            "attributes": {{
                "extraVariables": "comma-separated list if present",
                "compat": "compatibility level if present",
                "theme": "theme name if present",
                "style": "style name if present"
            }},
            "control_tags": {{
                "gotos": [{{"target": "Q10", "condition": "...", "position": "after_Q9"}}],
                "exits": [{{"action": "...", "condition": "...", "position": "after_Q20", "redirect_url": "..."}}],
                "finishes": [{{"condition": "...", "position": "after_Q25", "message": "..."}}],
                "terms": [{{"label": "Term_Q1", "condition": "...", "message": "...", "screen": "...", "position": "after_Q5"}}],
                "suspends": [{{"position": "after_Q1", "condition": "..."}}],
                "blocks": [{{"label": "Block1", "title": "...", "randomize": "...", "cond": "...", "position": "before_Q10"}}],
                "conditions": [{{"label": "Cond1", "logic": "...", "usage": ["Q10", "Q11"]}}],
                "alerts": [{{"condition": "...", "recipients": ["email@example.com"], "subject": "...", "message": "...", "position": "after_Q15"}}],
                "loops": [{{"label": "Loop1", "rows": [...], "questions": [...], "randomize": "...", "position": "after_Q8"}}],
                "inserts": [{{"target": "Q10", "source": "Q5", "condition": "...", "position": "after_Q9"}}],
                "pipes": [{{"condition": "...", "content": "...", "format": "...", "position": "in_Q10"}}],
                "datasources": [{{"name": "...", "file_path": "...", "column_mappings": {{...}}, "position": "..."}}],
                "defines": [{{"name": "...", "answer_list": {{...}}, "usage": [...], "position": "..."}}],
                "html": [{{"content": "...", "where": "...", "condition": "...", "position": "..."}}],
                "notes": [{{"content": "...", "placement": "...", "condition": "...", "position": "..."}}]
            }},
            "python_variables": {{
                "global": [{{"name": "...", "source": "extraVariables|URL|exec_init", "value": "..."}}],
                "local": [{{"name": "...", "scope": "question_id", "code": "..."}}],
                "persistent": [{{"name": "...", "code": "..."}}]
            }},
            "styling": {{
                "styles": [{{"name": "...", "content": "...", "where": "survey|report|execute", "condition": "..."}}],
                "stylevars": [{{"name": "...", "value": "...", "type": "...", "usage": "..."}}],
                "themes": [{{"name": "...", "content": "...", "attributes": {{...}}}}],
                "themevars": [{{"name": "...", "value": "...", "scope": "..."}}],
                "css_files": [{{"path": "path/to/file.css", "location": "...", "type": "..."}}],
                "js_files": [{{"path": "path/to/file.js", "location": "...", "type": "..."}}],
                "static_files": [{{"path": "path/to/file", "type": "image|font|other", "location": "..."}}],
                "less_css": {{"enabled": true|false, "content": "...", "variables": [...]}},
                "google_fonts": [{{"family": "...", "url": "...", "weights": [...]}}]
            }},
            "participant_sources": {{
                "samplesources": [{{"name": "...", "type": "URL|panel|API|other", "configuration": {{...}}, "validation": "...", "variables": [...]}}],
                "samplesources_multiple": [{{"sources": [...], "priority": "...", "conditions": "..."}}],
                "url_variables": [{{"name": "...", "values": [...], "validation": "...", "default": "..."}}],
                "locking": {{"enabled": true|false, "variable": "...", "condition": "...", "message": "..."}},
                "access_control": {{"one_time": true|false, "restrictions": [...], "conditions": "..."}},
                "completion_actions": [{{"type": "redirect|message|exit", "url": "...", "message": "...", "condition": "..."}}]
            }},
            "survey_attributes": {{
                "project_settings": {{
                    "extraVariables": "...",
                    "compat": "...",
                    "survey_mode": "live|test|preview",
                    "language_settings": {{"default": "...", "multi_language": true|false, "languages": [...]}}
                }},
                "display_settings": {{
                    "theme": {{"name": "...", "configuration": {{...}}}},
                    "style": {{"name": "...", "configuration": {{...}}}},
                    "mobile_optimization": {{...}},
                    "browser_settings": {{...}},
                    "display_mode": "..."
                }},
                "device_settings": {{
                    "mobile": {{...}},
                    "tablet": {{...}},
                    "desktop": {{...}},
                    "device_detection": {{...}}
                }},
                "field_settings": {{
                    "default_size": "...",
                    "default_validation": "...",
                    "default_styling": {{...}}
                }},
                "question_settings": {{
                    "default_attributes": {{...}},
                    "default_styling": {{...}},
                    "default_behavior": {{...}}
                }},
                "system_language_resources": {{
                    "error_messages": {{...}},
                    "tooltips": {{...}},
                    "button_text": {{...}},
                    "window_titles": {{...}},
                    "support_links": {{...}}
                }}
            }},
            "datasources": [
                {{
                    "name": "...",
                    "file_path": "path/to/file.txt",
                    "type": "tab_delimited|csv|other",
                    "column_mappings": {{...}},
                    "refresh_settings": {{...}},
                    "validation": "..."
                }}
            ],
            "defines": [
                {{
                    "name": "...",
                    "answer_list": {{"rows": [...], "columns": [...], "choices": [...]}},
                    "attributes": {{...}},
                    "usage": ["Q1", "Q2"]
                }}
            ],
            "html_elements": [
                {{
                    "content": "...",
                    "where": "survey|report|execute",
                    "condition": "...",
                    "formatting": "..."
                }}
            ],
            "notes": [
                {{
                    "content": "...",
                    "placement": "...",
                    "condition": "...",
                    "formatting": "..."
                }}
            ],
            "macros": [
                {{
                    "name": "...",
                    "parameters": [...],
                    "content": "...",
                    "attributes": {{...}},
                    "usage": [{{"location": "...", "parameters": [...]}}]
                }}
            ],
            "mutators": [
                {{
                    "function_name": "...",
                    "parameters": [...],
                    "logic": "...",
                    "structure_modifications": [{{"type": "create_question|modify_question|create_row", "details": {{...}}}}]
                }}
            ],
            "database_integration": {{
                "adb": {{
                    "enabled": true|false,
                    "configuration": {{...}},
                    "data_sources": [...],
                    "integration_patterns": [...]
                }},
                "auxiliary_database": {{
                    "enabled": true|false,
                    "configuration": {{...}},
                    "read_operations": [...],
                    "write_operations": [...]
                }},
                "data_loading": {{
                    "loadData_calls": [{{"source": "...", "condition": "..."}}],
                    "loadRecord_calls": [{{"source": "...", "condition": "..."}}]
                }}
            }},
            "transformation_system": {{
                "scripts": [
                    {{
                        "name": "...",
                        "content": "...",
                        "triggers": [{{"type": "completion|event|condition", "details": {{...}}}}]
                    }}
                ],
                "transformations": [
                    {{
                        "type": "content_modification|structure_change|data_transformation",
                        "logic": "...",
                        "conditions": "..."
                    }}
                ]
            }}
        }},
        "questions": [
            {{
                "id": "Q1",
                "text": "The actual question text",
                "type": "radio|checkbox|grid|select|number|float|text|textarea|virtual|logic|autosuggest|button_select|button_rating|card_rating",
                "cdata": false,
                "where": "survey|report|execute",
                "variable_scope": "global|local|persistent",
                "options": [
                    {{
                        "value": "1",
                        "text": "Option text",
                        "special_logic": "IF X THEN Y",
                        "termination_logic": "TERMINATE IF P1.r3",
                        "skip_logic": "SKIP TO Q10 IF SELECTED",
                        "exclusive": false,
                        "open": false,
                        "cond": "conditional_display_logic"
                    }}
                ],
                "conditions": "Any conditional logic",
                "section": "Section name if present",
                "instructions": ["Instruction 1", "Instruction 2", ...],
                "attributes": {{
                    "optional": "0|1",
                    "unique": "0|1",
                    "verify": "verification_rule",
                    "size": "input_size",
                    "atleast": "minimum_selections",
                    "atmost": "maximum_selections",
                    "shuffle": "rows|columns|both|none",
                    "randomize": "randomization_pattern",
                    "cond": "conditional_display_logic",
                    "di": "display_instruction",
                    "ss": "single_select_mode",
                    "adim": "report_dimension",
                    "virtual": "virtual_calculation"
                }},
                "cell_attributes": {{
                    "rows": [{{"label": "r1", "open": "0|1", "exclusive": "0|1", "cond": "...", "value": "..."}}],
                    "columns": [{{"label": "c1", "open": "0|1", "exclusive": "0|1", "cond": "...", "value": "..."}}],
                    "choices": [{{"label": "ch1", "open": "0|1", "exclusive": "0|1", "cond": "...", "value": "..."}}]
                }},
                "exec_blocks": [
                    {{
                        "when": "init|virtualInit|virtual|none",
                        "code": "Python code here",
                        "purpose": "variable_assignment|calculation|logic"
                    }}
                ],
                "virtual": "calculation_logic_for_virtual_questions",
                "validation": {{
                    "rules": [{{"type": "range|required|custom", "min": 1, "max": 100, "message": "..."}}],
                    "messages": ["Error message 1", "Error message 2"],
                    "conditions": "IF condition THEN validate"
                }},
                "nets": [
                    {{
                        "label": "Net1",
                        "type": "sum|average|count",
                        "rows": ["r1", "r2"],
                        "columns": [],
                        "choices": []
                    }}
                ],
                "quota_info": {{
                    "target": "Target audience or variable name",
                    "limits": ["Target limits from text like N=1000, etc."],
                    "qualification_logic": "ACTUAL extracted conditions - NEVER use generic examples",
                    "termination_conditions": "Termination rules if any"
                }},
                "table_data": {{
                    "headers": [...],
                    "rows": [...]
                }},
                "subquestions": [...],
                "recode_info": {{
                    "source_question": "Q1",
                    "logic": "Recode logic",
                    "value_mappings": [...]
                }},
                "hidden_info": {{
                    "purpose": "Purpose of hidden question",
                    "associated_questions": [...]
                }},
                "preload_info": {{
                    "source": "URL parameter or external source",
                    "default_value": "Default value if not provided",
                    "validation": "Validation rules for preload values"
                }},
                "rotation_info": {{
                    "type": "rows|columns|both",
                    "randomization": true|false
                }},
                "images": [
                    {{
                        "source": "path/to/image.jpg",
                        "alt": "alt text",
                        "title": "image title",
                        "description": "image description",
                        "caption": "image caption",
                        "width": "width in pixels or percentage",
                        "height": "height in pixels or percentage",
                        "alignment": "left|right|center|inline",
                        "position": "positioning information",
                        "format": "jpg|png|gif|svg|webp",
                        "placeholder": false
                    }}
                ],
                "vector_logic": false,
                "randomization": {{
                    "type": "none|rows|columns|both|balanced|latin_square|block",
                    "shuffle": "rows|columns|both|none",
                    "condition": "conditional randomization",
                    "rotation": "rotation pattern"
                }}
            }}
        ]
    }}

    ENHANCED QUESTION STRUCTURE EXAMPLES:
    
    **Multi-Select Question with Instructions:**
    {{
        "id": "S4",
        "text": "Which of the following best describes your racial background?",
        "type": "checkbox",
        "options": [
            {{"value": "1", "text": "White or Caucasian"}},
            {{"value": "2", "text": "Hispanic or Latino"}},
            {{"value": "3", "text": "Black or African American"}},
            {{"value": "4", "text": "Asian or Asian American"}},
            {{"value": "5", "text": "Native Hawaiian or Other Pacific Islander"}},
            {{"value": "6", "text": "Indigenous American or Indigenous Alaskan"}},
            {{"value": "7", "text": "Middle Eastern or North African"}},
            {{"value": "8", "text": "Other"}},
            {{"value": "9", "text": "Prefer not to answer", "exclusive": true}}
        ],
        "conditions": {{
            "display_logic": "Always show",
            "skip_logic": "None"
        }},
        "instructions": [
            "MULTI SELECT - Allow multiple selections",
            "Option 9 is EXCLUSIVE - If selected, deselect all others",
            "Use checkbox question type with atleast='1'"
        ],
        "comments": [
            "Please select all that apply."
        ]
    }}
    
    **Conditional Question with Complex Logic:**
    {{
        "id": "S9",
        "text": "As an independent insurance agent, what best describes your role?",
        "type": "radio",
        "options": [
            {{"value": "1", "text": "Principal or owner of the company"}},
            {{"value": "2", "text": "Producer or agent at the company"}},
            {{"value": "3", "text": "Customer service representative at the company"}},
            {{"value": "4", "text": "None of the above", "termination_logic": "IF P1.r3 THEN TERMINATE"}}
        ],
        "conditions": {{
            "display_logic": "ASK IF S8.r1",
            "skip_logic": "Skip if S8.r2, S8.r3, etc.",
            "nested_logic": "IF S8.r1 THEN SHOW S9"
        }},
        "instructions": [
            "SINGLE SELECT - Only one option allowed",
            "Option 4 has termination logic: IF P1.r3 THEN TERMINATE",
            "Use radio question type",
            "Add termination logic for option 4 when P1.r3 is selected"
        ],
        "comments": []
    }}
    
    **Quota Questions (S100):**
    {{
        "id": "S100",
        "text": "OVERALL QUALIFICATIONS",
        "type": "checkbox",
        "options": ["U.S. WORKFORCE", "TECH WORKFORCE", "BUSINESS DECISION-MAKER", "NOT QUALIFIED"],
        "quota_info": {{
            "target": "S100",
            "limits": ["N=1,600", "N=600", "N=200", "TERMINATE"],
            "qualification_logic": "EXTRACT THE ACTUAL LOGIC FROM THE QUESTION TEXT",
            "termination_conditions": "EXTRACT THE ACTUAL TERMINATION CONDITIONS FROM THE QUESTION"
        }},
        "conditions": {{
            "display_logic": "Always show",
            "skip_logic": "None"
        }},
        "instructions": ["EXTRACT ACTUAL LOGIC FROM QUESTION TEXT", "DO NOT USE GENERIC EXAMPLES"],
        "comments": []
    }}
    
    **Preload Variables (P1):**
    {{
        "id": "P1",
        "text": "HIDDEN QUESTION - PRELOADS - Set from URL parameter",
        "type": "preload",
        "options": ["Preload 1", "Preload 2"],
        "preload_info": {{
            "source": "EXTRACT THE ACTUAL SOURCE FROM THE QUESTION TEXT",
            "default_value": "EXTRACT THE ACTUAL DEFAULT VALUE FROM THE QUESTION",
            "validation": "EXTRACT THE ACTUAL VALIDATION RULES FROM THE QUESTION"
        }},
        "conditions": {{
            "display_logic": "Hidden from respondent",
            "skip_logic": "None"
        }},
        "instructions": ["Set value from URL parameter 'preload'", "Terminate if preload not passed"],
        "comments": []
    }}

    ENHANCED EXTRACTION RULES:
    
    **STRUCTURED DATA EXTRACTION:**
    1. **PRESERVE EXACT ORDER** of questions - never skip or reorder
    2. **EXTRACT ALL REPEATED QUESTIONS** - each occurrence gets its own entry
    3. Extract ALL response options, scales, and table data
    4. Maintain original formatting and structure
    
    **CONDITIONS EXTRACTION:**
    5. **Display Logic**: Extract [ASK IF X], [SHOW IF X], [ONLY SHOW IF X] patterns
    6. **Skip Logic**: Extract [SKIP TO X], [SKIP IF X] patterns
    7. **Nested Logic**: Capture complex IF-THEN-ELSE structures
    8. **Multiple Conditions**: Extract AND/OR combinations correctly
    9. **Reference Logic**: Identify question dependencies (IF S8/1 THEN S9)
    
    **INSTRUCTIONS EXTRACTION:**
    10. **Question Type**: Extract [MULTI SELECT], [SINGLE SELECT], [RADIO], [CHECKBOX]
    11. **Special Options**: Extract [EXCLUSIVE], [TERMINATE IF X], [HIDDEN]
    12. **Programming Notes**: Extract technical implementation details
    13. **Validation Rules**: Extract constraints, limits, and requirements
    14. **Termination Logic**: Extract termination conditions for specific options
    
    **COMMENTS EXTRACTION:**
    15. **Respondent Guidance**: Extract "Please select all that apply", "Choose one", etc.
    16. **Clarification Text**: Extract explanatory text for respondents
    17. **Instructions Text**: Extract any text that guides respondent behavior
    18. **ONLY include actual respondent guidance in comments**
    
    **SPECIAL QUESTION TYPES:**
    19. **Quota Questions**: Extract S100 series with complete qualification logic
    20. **Preload Variables**: Extract P1 series with source and validation
    21. **Hidden Questions**: Extract internal logic and tracking variables
    22. **Recode Questions**: Extract computed variables and transformations
    
    **CRITICAL EXTRACTION RULES:**
    23. **CRITICAL**: Extract ALL conditions with proper display/skip logic
    24. **CRITICAL**: Extract ALL instructions for accurate XML generation
    25. **CRITICAL**: Extract ALL comments for respondent guidance
    26. **CRITICAL**: Extract ALL termination logic and conditions
    27. **CRITICAL**: Extract ALL exclusive options and their behavior
    28. **CRITICAL**: Extract ALL nested conditional logic correctly
    29. **CRITICAL**: NEVER skip repeated questions - extract ALL occurrences
    30. **CRITICAL**: Continue processing all content after section headers

    **CRITICAL QUESTION TEXT EXTRACTION RULES:**
    - **QUESTION TEXT**: Extract the ACTUAL question text that respondents will see, NOT the question ID or label
    - **QUESTION ID**: Extract the question identifier (e.g., "TEMP_AGE", "UC3", "Q1")
    - **QUESTION LABEL**: Extract the descriptive label in parentheses (e.g., "(Initial Age)", "(HH Industry)")
    
    **EXAMPLES OF CORRECT EXTRACTION:**
    
    Input: "(ASK ALL) TEMP_AGE. (Initial Age) What is your age? (OE NUMERIC, RANGE 1-99)"
    Correct Output:
    ```json
    {{
        "id": "TEMP_AGE",
        "text": "What is your age?",
        "type": "number",
        "instructions": ["Please enter your age in years."]
    }}
    ```
    
    Input: "(ASK ALL) UC3. (HH Industry) Do you or does anyone in your household work for any of the following: (MULTIPLE SELECT, RANDOMIZE)"
    Correct Output:
    ```json
    {{
        "id": "UC3", 
        "text": "Do you or does anyone in your household work for any of the following:",
        "type": "checkbox",
        "instructions": ["Please select all that apply"]
    }}
    ```
    
    **WRONG EXTRACTION (DO NOT DO THIS):**
    ```json
    {{
        "id": "TEMP_AGE",
        "text": "Initial Age",  // WRONG - this is the label, not the question
        "type": "number"
    }}
    ```

    **CRITICAL OPTION-LEVEL LOGIC EXTRACTION:**
    - **EVERY OPTION** must be analyzed for embedded logic, conditions, and special behavior
    - **EXTRACT ALL LOGIC** associated with each option, even if it seems minor
    - **CAPTURE TERMINATION LOGIC**: "TERMINATE IMMEDIATELY", "DISQUALIFY IF SELECTED"
    - **CAPTURE SKIP LOGIC**: "SKIP TO Q10", "GO TO SECTION 2"
    - **CAPTURE CONDITIONAL DISPLAY**: "SHOW Q15 ONLY IF SELECTED", "DISPLAY IF Q1 = 1"
    - **CAPTURE VALIDATION RULES**: "MUST BE 18+", "RANGE 1-99", "REQUIRED IF SELECTED"
    - **CAPTURE EXCLUSIVE BEHAVIOR**: "NONE OF THE ABOVE", "OTHER (SPECIFY)"
    - **CAPTURE DEPENDENCY LOGIC**: "ONLY IF Q1 = YES", "IF SELECTED, SHOW Q10"
    - **CAPTURE SPECIAL INSTRUCTIONS**: "DO NOT READ OUT LOUD", "INTERVIEWER NOTE"
    
    **EXAMPLES OF OPTION-LEVEL LOGIC EXTRACTION:**
    
    Input: "1. The production, distribution, or retailing of video games (TERMINATE IMMEDIATELY)"
    Correct Output:
    ```json
    {{
        "value": "1",
        "text": "The production, distribution, or retailing of video games",
        "termination_logic": "TERMINATE IMMEDIATELY",
        "skip_logic": null,
        "conditional_display": null,
        "validation_rules": null,
        "exclusive_behavior": null,
        "dependency_logic": null,
        "special_instructions": null
    }}
    ```
    
    Input: "5. Mobile app developer or publisher (TERMINATE IMMEDIATELY)"
    Correct Output:
    ```json
    {{
        "value": "5",
        "text": "Mobile app developer or publisher",
        "termination_logic": "TERMINATE IMMEDIATELY",
        "skip_logic": null,
        "conditional_display": null,
        "validation_rules": null,
        "exclusive_behavior": null,
        "dependency_logic": null,
        "special_instructions": null
    }}
    ```
    
    Input: "8. Consumer electronics manufacturer"
    Correct Output:
    ```json
    {{
        "value": "8",
        "text": "Consumer electronics manufacturer",
        "termination_logic": null,
        "skip_logic": null,
        "conditional_display": null,
        "validation_rules": null,
        "exclusive_behavior": null,
        "dependency_logic": null,
        "special_instructions": null
    }}
    ```

    **ENHANCED OUTPUT FORMAT:**
    Each question must include these structured fields:
    
    ```json
    {{
        "id": "QuestionID",
        "text": "Question text",
        "type": "question_type",
        "options": [
            {{
                "value": "1", 
                "text": "Option text", 
                "termination_logic": "TERMINATE IMMEDIATELY IF SELECTED",
                "skip_logic": "SKIP TO Q10 IF SELECTED",
                "conditional_display": "SHOW Q15 ONLY IF SELECTED",
                "validation_rules": "MUST BE 18+",
                "exclusive_behavior": "DESELECTS OTHER OPTIONS",
                "dependency_logic": "ONLY IF Q1 = YES",
                "special_instructions": "DO NOT READ OUT LOUD"
            }},
            {{
                "value": "2", 
                "text": "Option text",
                "termination_logic": null,
                "skip_logic": null,
                "conditional_display": null,
                "validation_rules": null,
                "exclusive_behavior": null,
                "dependency_logic": null,
                "special_instructions": null
            }}
        ],
        "conditions": {{
            "display_logic": "ASK IF X.r1",
            "skip_logic": "SKIP TO Y IF Z",
            "nested_logic": "IF A AND B THEN SHOW"
        }},
        "instructions": [
            "PROGRAMMING NOTE: Use checkbox with atleast='1'",
            "Option 4 has termination logic: IF P1.r3 THEN TERMINATE"
        ],
        "comments": [
            "Please select all that apply.",
            "Choose the option that best describes you."
        ]
    }}
    ```

    Chunk text to analyze:
    """
    
    if client is None:
        raise ValueError("OpenAI client is not available. Please check OPENAI_API_KEY environment variable.")
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert survey document parser specialized in Forsta Surveys XML format. Extract complete survey structures including: all question types (radio, checkbox, text, textarea, number, float, select, grid, virtual, logic), table data, response options, scales, metadata, recodes, hidden questions, special processing rules, CDATA sections, Python variables (global/local/persistent), exec blocks with when attributes, control tags (goto, exit, finish, term, suspend, block, condition, alert, loop, insert, pipe), validation rules, net aggregations, complete attribute extraction, vector logic patterns, advanced randomization patterns, enhanced image metadata, complete styling system, complete participant source configuration, enhanced survey attributes, datasource tags, define tags, HTML/note tags, macro system, mutator functions, database integration, transformation system, and ALL instructions or programming notes. Be thorough and precise. For the 'type' field, use only the Decipher XML types as per the mapping provided. Do not use descriptive labels. Do not miss any instructions or interviewer notes. You must extract every question, preload, recode, hidden, derived, virtual, or logic-referenced variable. Do not skip any item that could be used as a question or variable, even if it is not directly shown to the respondent. Extract survey-level metadata including XML declaration, survey attributes (all categories), control tags with positions, Python variables by scope, complete styling system, complete participant sources, datasources, defines, HTML/note elements, macros, mutators, database integration, and transformation system. Extract vector logic operations, advanced randomization patterns, and complete image metadata. If there is any uncertainty, err on the side of inclusion and treat it as a question or variable. Extract all variable names, programming notes, and logic-referenced items as questions or variables in the output."},
                {"role": "user", "content": prompt + chunk_text}
            ],
            temperature=0.1,
            max_tokens=16384,  # Correct maximum for GPT-4o
            response_format={"type": "json_object"}
        )
        
        # Get the response content
        response_content = response.choices[0].message.content
        
        # Try to parse the JSON response
        try:
            # Clean the response content
            cleaned_content = response_content.strip()
            if cleaned_content.startswith('\ufeff'):
                cleaned_content = cleaned_content[1:]
            
            # Enhanced JSON cleaning to handle common LLM response issues
            cleaned_content = _clean_llm_response_for_json(cleaned_content)
            
            # Parse JSON
            result = json.loads(cleaned_content)
            
            # Validate structure - support both old format (questions array) and new format (survey + questions)
            if not isinstance(result, dict):
                print(f"Invalid response format from chunk {chunk_index + 1}: Response is not a dictionary")
                return None
            
            # Check for new format with survey metadata
            if 'survey' in result:
                # New format with survey-level metadata
                if 'questions' not in result:
                    print(f"Invalid response format from chunk {chunk_index + 1}: Missing questions array in survey structure")
                    return None
                if not isinstance(result['questions'], list):
                    print(f"Invalid response format from chunk {chunk_index + 1}: Questions field is not a list")
                    return None
            elif 'questions' in result:
                # Old format - just questions array (backward compatibility)
                if not isinstance(result['questions'], list):
                    print(f"Invalid response format from chunk {chunk_index + 1}: Questions field is not a list")
                    return None
                # Convert to new format for consistency
                result = {'survey': {}, 'questions': result['questions']}
            else:
                print(f"Invalid response format from chunk {chunk_index + 1}: Missing both 'survey' and 'questions' fields")
                return None
                
            # Validate each question
            for i, question in enumerate(result['questions']):
                if not isinstance(question, dict):
                    print(f"Invalid question format at index {i} in chunk {chunk_index + 1}")
                    continue
                    
                required_fields = ['id', 'text']
                for field in required_fields:
                    if field not in question:
                        print(f"Missing required field '{field}' in question {i} of chunk {chunk_index + 1}")
            
            return result
            
        except json.JSONDecodeError as e:
            print(f"Error parsing response from chunk {chunk_index + 1}: {str(e)}")
            print(f"Response content preview: {response_content[:200]}...")
            print(f"Cleaned content preview: {cleaned_content[:200]}...")
            
            # ROBUST PARTIAL EXTRACTION - Extract ALL questions regardless of JSON format
            print(f"Attempting robust partial extraction for chunk {chunk_index + 1}...")
            
            try:
                import re
                extracted_questions = []
                
                # METHOD 1: Extract complete question objects using flexible regex
                complete_question_pattern = r'\{"id":"([^"]+)"[^}]*"text":"([^"]+)"[^}]*"type":"([^"]+)"[^}]*\}'
                complete_matches = re.findall(complete_question_pattern, response_content)
                
                if complete_matches:
                    print(f"Method 1: Found {len(complete_matches)} complete question objects")
                    for match in complete_matches:
                        question_obj = {
                            "id": match[0],
                            "text": match[1],
                            "type": match[2],
                            "options": [],
                            "extraction_method": "complete_object"
                        }
                        extracted_questions.append(question_obj)
                
                # METHOD 2: Extract individual fields and reconstruct questions
                print(f"Method 2: Extracting individual fields...")
                
                # Extract all question IDs
                id_pattern = r'"id"\s*:\s*"([^"]+)"'
                all_ids = re.findall(id_pattern, response_content)
                print(f"  Found {len(all_ids)} question IDs: {all_ids[:10]}{'...' if len(all_ids) > 10 else ''}")
                
                # Extract all question texts (handle very long text)
                text_pattern = r'"text"\s*:\s*"([^"]{1,200})'
                all_texts = re.findall(text_pattern, response_content)
                print(f"  Found {len(all_texts)} question texts")
                
                # Extract all question types
                type_pattern = r'"type"\s*:\s*"([^"]+)"'
                all_types = re.findall(type_pattern, response_content)
                print(f"  Found {len(all_types)} question types: {all_types[:10]}{'...' if len(all_types) > 10 else ''}")
                
                # Extract all options arrays
                options_pattern = r'"options"\s*:\s*\[([^\]]*)\]'
                all_options = re.findall(options_pattern, response_content)
                print(f"  Found {len(all_options)} options arrays")
                
                # METHOD 3: Smart question reconstruction
                max_questions = max(len(all_ids), len(all_texts), len(all_types))
                print(f"  Maximum potential questions: {max_questions}")
                
                # Reconstruct questions from extracted fields
                for i in range(max_questions):
                    question_obj = {}
                    
                    # Add ID if available
                    if i < len(all_ids):
                        question_obj["id"] = all_ids[i]
                    
                    # Add text if available
                    if i < len(all_texts):
                        question_obj["text"] = all_texts[i] + ("..." if len(all_texts[i]) == 200 else "")
                    
                    # Add type if available
                    if i < len(all_types):
                        question_obj["type"] = all_types[i]
                    
                    # Add options if available
                    if i < len(all_options):
                        options_str = all_options[i]
                        if options_str.strip():
                            option_pattern = r'"([^"]+)"'
                            options = re.findall(option_pattern, options_str)
                            question_obj["options"] = options
                        else:
                            question_obj["options"] = []
                    else:
                        question_obj["options"] = []
                    
                    # Add extraction metadata
                    question_obj["extraction_method"] = "field_reconstruction"
                    question_obj["extraction_note"] = f"Reconstructed from chunk {chunk_index + 1}"
                    
                    # Only add if we have at least ID and text
                    if "id" in question_obj and "text" in question_obj:
                        extracted_questions.append(question_obj)
                
                # METHOD 4: Look for any remaining question-like patterns
                remaining_pattern = r'"([A-Z][0-9A-Z_]*)"\s*:\s*"([^"]{10,})'
                remaining_matches = re.findall(remaining_pattern, response_content)
                
                if remaining_matches:
                    print(f"Method 4: Found {len(remaining_matches)} additional potential questions")
                    for match in remaining_matches:
                        # Check if this looks like a question ID
                        if re.match(r'^[A-Z][0-9A-Z_]*$', match[0]):
                            question_obj = {
                                "id": match[0],
                                "text": match[1][:100] + ("..." if len(match[1]) > 100 else ""),
                                "type": "text",  # Default type
                                "options": [],
                                "extraction_method": "pattern_matching",
                                "extraction_note": f"Found via pattern matching in chunk {chunk_index + 1}"
                            }
                            extracted_questions.append(question_obj)
                
                # Remove duplicates based on ID
                unique_questions = []
                seen_ids = set()
                for q in extracted_questions:
                    if q["id"] not in seen_ids:
                        unique_questions.append(q)
                        seen_ids.add(q["id"])
                
                print(f"Total unique questions extracted: {len(unique_questions)}")
                
                if unique_questions:
                    # Sort questions by ID for consistency
                    unique_questions.sort(key=lambda x: x["id"])
                    
                    partial_result = {"questions": unique_questions}
                    print(f"✅ SUCCESS: Extracted {len(unique_questions)} questions using robust partial extraction")
                    return partial_result
                else:
                    print(f"❌ No questions could be extracted from chunk {chunk_index + 1}")
                    return None
                    
            except Exception as partial_e:
                print(f"Robust partial extraction failed: {str(partial_e)}")
                import traceback
                traceback.print_exc()
                return None
            
    except Exception as e:
        print(f"Error processing chunk {chunk_index + 1} with GPT-4: {str(e)}")
        return None


# ============================================================================
# Main Processing Function
# ============================================================================

def process_text_with_batch_gpt4(text, use_parallel=False):
    """
    Process large text using batch processing with GPT-4. Option for parallel processing.
    
    Args:
        text: Text content to process
        use_parallel: Whether to use parallel processing for large documents
        
    Returns:
        Dictionary with 'survey' and 'questions' keys
    """
    try:
        # Simple size check - use character count for reliability
        if len(text) <= 20000:
            # Small enough to process in one go
            result = process_chunk_with_gpt4(text, 0, 1)
            if result and 'survey' not in result:
                result = {'survey': {}, 'questions': result.get('questions', [])}
            return result
        
        # Split into chunks
        max_chunk_size = 20000
        chunks = split_text_into_chunks(text, max_chars=max_chunk_size)
        
        print(f"Large document detected! Processing in {len(chunks)} chunks.")
        print(f"Chunk size limit: {max_chunk_size:,} characters")
        
        # Show chunk sizes
        for i, chunk in enumerate(chunks):
            chunk_chars = len(chunk)
            chunk_tokens = count_tokens_estimate(chunk)
            print(f"Chunk {i+1}: {chunk_chars:,} characters (~{chunk_tokens:,} tokens)")
        
        all_questions = []
        survey_metadata = {
            'xml_declaration': {},
            'attributes': {},
            'control_tags': {},
            'python_variables': {'global': [], 'local': [], 'persistent': []},
            'styling': {
                'styles': [],
                'stylevars': [],
                'themes': [],
                'themevars': [],
                'css_files': [],
                'js_files': [],
                'static_files': [],
                'less_css': {},
                'google_fonts': []
            },
            'participant_sources': {
                'samplesources': [],
                'samplesources_multiple': [],
                'url_variables': [],
                'locking': {},
                'access_control': {},
                'completion_actions': []
            },
            'survey_attributes': {
                'project_settings': {},
                'display_settings': {},
                'device_settings': {},
                'field_settings': {},
                'question_settings': {},
                'system_language_resources': {}
            },
            'datasources': [],
            'defines': [],
            'html_elements': [],
            'notes': [],
            'macros': [],
            'mutators': [],
            'database_integration': {
                'adb': {},
                'auxiliary_database': {},
                'data_loading': {}
            },
            'transformation_system': {
                'scripts': [],
                'transformations': []
            }
        }
        successful_chunks = 0
        failed_chunks = []
        
        if use_parallel and len(chunks) > 2:
            # Parallel processing for large surveys
            print("Processing chunks in parallel...")
            
            def process_chunk_wrapper(chunk_data):
                chunk_idx, chunk_text = chunk_data
                return chunk_idx, process_chunk_with_gpt4(chunk_text, chunk_idx, len(chunks))
            
            # Process chunks in parallel with limited workers
            with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
                chunk_data = [(i, chunk) for i, chunk in enumerate(chunks)]
                future_to_chunk = {executor.submit(process_chunk_wrapper, data): data for data in chunk_data}
                
                # Collect results in a list to maintain order
                chunk_results = [None] * len(chunks)
                completed = 0
                
                for future in concurrent.futures.as_completed(future_to_chunk):
                    chunk_idx, chunk_result = future.result()
                    completed += 1
                    
                    # Store result in correct position
                    chunk_results[chunk_idx] = chunk_result
                    
                    print(f"Completed chunk {chunk_idx + 1} of {len(chunks)}...")
                    
                    if chunk_result and 'questions' in chunk_result:
                        print(f"Chunk {chunk_idx + 1}: Extracted {len(chunk_result['questions'])} questions")
                    else:
                        print(f"Chunk {chunk_idx + 1}: No questions extracted")
                
                # Process results in correct order
                for i, chunk_result in enumerate(chunk_results):
                    if chunk_result:
                        # Extract questions
                        if 'questions' in chunk_result:
                            all_questions.extend(chunk_result['questions'])
                        
                        # Merge survey-level metadata (first chunk wins for metadata, but merge arrays)
                        if 'survey' in chunk_result and i == 0:
                            chunk_survey = chunk_result['survey']
                            for key, value in chunk_survey.items():
                                if key in survey_metadata:
                                    if isinstance(value, dict) and isinstance(survey_metadata[key], dict):
                                        survey_metadata[key].update(value)
                                    elif isinstance(value, list) and isinstance(survey_metadata[key], list):
                                        survey_metadata[key].extend(value)
                                    else:
                                        survey_metadata[key] = value
                                else:
                                    survey_metadata[key] = value
                        
                        successful_chunks += 1
                    else:
                        failed_chunks.append(i + 1)
        else:
            # Sequential processing
            print("Processing chunks sequentially...")
            for i, chunk in enumerate(chunks):
                print(f"Processing chunk {i+1} of {len(chunks)}...")
                
                # Reduced delay to avoid rate limits - GPT-4o can handle faster requests
                if i > 0:
                    time.sleep(5)  # Reduced from 30 to 5 seconds
                
                chunk_result = process_chunk_with_gpt4(chunk, i, len(chunks))
                
                if chunk_result:
                    # Extract questions
                    if 'questions' in chunk_result:
                        all_questions.extend(chunk_result['questions'])
                    
                    # Merge survey-level metadata (first chunk wins)
                    if 'survey' in chunk_result and i == 0:
                        chunk_survey = chunk_result['survey']
                        for key, value in chunk_survey.items():
                            if key in survey_metadata:
                                if isinstance(value, dict) and isinstance(survey_metadata[key], dict):
                                    survey_metadata[key].update(value)
                                elif isinstance(value, list) and isinstance(survey_metadata[key], list):
                                    survey_metadata[key].extend(value)
                                else:
                                    survey_metadata[key] = value
                            else:
                                survey_metadata[key] = value
                    
                    successful_chunks += 1
                else:
                    failed_chunks.append(i + 1)
        
        # Final result
        result = {
            'survey': survey_metadata,
            'questions': all_questions
        }
        
        print(f"\n✓ Processing complete!")
        print(f"  - Successful chunks: {successful_chunks}/{len(chunks)}")
        print(f"  - Total questions extracted: {len(all_questions)}")
        if failed_chunks:
            print(f"  - Failed chunks: {failed_chunks}")
        
        return result
        
    except Exception as e:
        print(f"Error in batch processing: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


# ============================================================================
# Complete Workflow Function
# ============================================================================

def extract_document_to_json(file_path, output_path=None, use_parallel=False):
    """
    Complete workflow: Extract text from document and convert to JSON.
    For DOCX files, converts to PDF first before processing.
    
    Args:
        file_path: Path to the input document (PDF, DOCX, or TXT)
        output_path: Path to save the JSON output (optional)
        use_parallel: Whether to use parallel processing for large documents
        
    Returns:
        Dictionary with extracted survey structure and questions
    """
    print(f"Starting extraction from: {file_path}")
    print("-" * 60)
    
    # Step 0: Convert DOCX to PDF if needed
    temp_pdf_path = None
    if file_path.lower().endswith('.docx'):
        print("Step 0: Converting DOCX to PDF...")
        temp_pdf_path = convert_docx_to_pdf(file_path)
        if not temp_pdf_path:
            print("❌ Failed to convert DOCX to PDF")
            return None
        file_path = temp_pdf_path  # Use the converted PDF for processing
    
    # Step 1: Extract text from document
    print("\nStep 1: Extracting text from document...")
    text = extract_text_from_document(file_path)
    
    if not text:
        print("❌ Failed to extract text from document")
        return None
    
    print(f"✓ Extracted {len(text):,} characters")
    print(f"  Estimated tokens: ~{count_tokens_estimate(text):,}")
    
    # Step 2: Process text with GPT-4
    print("\nStep 2: Processing text with GPT-4...")
    use_parallel = use_parallel or len(text) > 100000  # Auto-enable for large docs
    
    result = process_text_with_batch_gpt4(text, use_parallel=use_parallel)
    
    if not result or 'questions' not in result:
        print("❌ Failed to extract questions from text")
        return None
    
    # Step 3: Save to JSON file
    if output_path:
        print(f"\nStep 3: Saving results to {output_path}...")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        print(f"✓ Results saved to {output_path}")
    else:
        # Generate default output path
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        output_path = f"{base_name}_extracted.json"
        print(f"\nStep 3: Saving results to {output_path}...")
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(result, f, indent=2, ensure_ascii=False)
        print(f"✓ Results saved to {output_path}")
    
    print("\n" + "=" * 60)
    print("EXTRACTION COMPLETE!")
    print(f"  - Input file: {file_path}")
    print(f"  - Output file: {output_path}")
    print(f"  - Questions extracted: {len(result['questions'])}")
    print("=" * 60)
    
    # Clean up temporary PDF file if it was created from DOCX conversion
    if temp_pdf_path and os.path.exists(temp_pdf_path):
        try:
            os.remove(temp_pdf_path)
            print(f"\n✓ Cleaned up temporary PDF file: {os.path.basename(temp_pdf_path)}")
        except Exception as e:
            print(f"\n⚠ Warning: Could not delete temporary PDF file: {str(e)}")
    
    return result


# ============================================================================
# Main Entry Point (for command-line usage)
# ============================================================================

if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python app.py <file_path> [output_path] [--parallel]")
        print("\nExamples:")
        print("  python app.py survey.pdf")
        print("  python app.py survey.docx output.json")
        print("  python app.py survey.txt output.json --parallel")
        sys.exit(1)
    
    file_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 and not sys.argv[2].startswith('--') else None
    use_parallel = '--parallel' in sys.argv
    
    if not os.path.exists(file_path):
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    
    result = extract_document_to_json(file_path, output_path, use_parallel)
    
    if result:
        print(f"\nSuccessfully extracted {len(result['questions'])} questions!")
    else:
        print("\nFailed to extract questions from document.")
        sys.exit(1)

